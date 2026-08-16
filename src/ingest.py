"""
Ingestion: turns raw 3GPP documents (PDF, DOCX, or .txt) into structure-aware
chunks.

Note on 3GPP's real distribution format: specs downloaded from
https://www.3gpp.org/ftp/Specs/latest/<Release>/<series>_series are .zip
files (e.g. 24501-id0.zip for TS 24.501). Each zip contains a single .doc or
.docx file — not a PDF. Unzip the archive and drop the .docx into
data/sample_3gpp/; this script reads .docx directly (see read_docx_file).
If you only have the older binary .doc format, convert it to .docx first
(e.g. open and re-save in Word, or use LibreOffice: `soffice --convert-to docx`).

Why structure-aware chunking matters for hallucination reduction:
3GPP specs use rigid hierarchical clause numbering (e.g. "5.5.1.2.3 Initial
registration accepted by the network"). If we split naively by character
count, we routinely cut a clause in half and hand the model incomplete
context, which is a direct cause of hallucinated or fabricated procedure
steps. Instead we detect clause headers with a regex and treat each clause
as an atomic unit, only sub-splitting a clause if it exceeds MAX_CHUNK_CHARS.
Every chunk keeps its clause number + title + source document as metadata,
which lets the generator cite exact clauses instead of vague paraphrases.
"""
import os
import re
import json
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import DATA_DIR, CHUNKS_PATH, INDEX_DIR, MAX_CHUNK_CHARS, CHUNK_OVERLAP_CHARS

# Matches lines like "5.5.1.2.3 Initial registration accepted by the network"
CLAUSE_HEADER_RE = re.compile(r"^(\d+(?:\.\d+){0,5})\s+(.+)$")

# Table-of-contents lines look identical to real clause headers (a clause
# number followed by a title) but end in a tab-separated page number, e.g.
# "5.5.1\tRegistration procedure\t50". Real in-body clause headers never end
# this way, so this is a reliable way to filter ToC noise out of the index
# without a separate "skip until body starts" heuristic.
TOC_ENTRY_RE = re.compile(r"\t\d+\s*$")

# Matches "3GPP TS 24.501 - Non-Access-Stratum ..." as the doc title line
DOC_TITLE_RE = re.compile(r"^3GPP\s+(TS|TR)\s+([\d.]+)\s*-\s*(.+)$")


def read_text_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def read_pdf_file(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx_file(path):
    """3GPP's real distribution format is .docx inside a .zip. python-docx
    reads paragraph text; tables (used for some annexes) are skipped for
    simplicity — extend this if you need annex table content indexed too."""
    import docx
    document = docx.Document(path)
    return "\n".join(p.text for p in document.paragraphs)


def split_into_clauses(raw_text, doc_id, doc_title):
    """Split a document's text into (clause_number, clause_title, body) tuples."""
    lines = raw_text.splitlines()
    clauses = []
    current = None

    for line in lines:
        if TOC_ENTRY_RE.search(line):
            continue  # skip table-of-contents lines entirely — not real clause content
        stripped = line.strip()
        m = CLAUSE_HEADER_RE.match(stripped)
        # Real clause titles start with either a capital letter ("Scope",
        # "References", "General") or a digit — many real 3GPP IE names
        # begin with a number, e.g. "5GS mobile identity", "5G-GUTI",
        # "4G-GUTI". This rejects only the actual false-positive pattern:
        # stray lines from the front-matter "version numbering convention"
        # boilerplate (e.g. "3 or greater indicates TSG approved document
        # under change control."), which start with a bare digit followed
        # by a LOWERCASE word — never a digit followed by an uppercase
        # word like "5GS" or "5G-GUTI". An earlier version of this check
        # used "not isupper()" instead of "islower()", which also (and
        # incorrectly) rejected any title starting with a digit — silently
        # dropping real clauses like "9.11.3.4 5GS mobile identity" from
        # the corpus entirely. That specific miss was traced end-to-end
        # from a hallucinated citation in generation, back through
        # retrieval, to this exact filtering bug.
        if m and m.group(2)[:1].islower():
            m = None
        if m and len(m.group(2)) < 120:  # header lines are short; avoids false positives on body text
            if current:
                clauses.append(current)
            current = {
                "clause_number": m.group(1),
                "clause_title": m.group(2).strip(),
                "body_lines": [],
            }
        else:
            if current is not None:
                current["body_lines"].append(line)
            # lines before the first detected clause header are dropped (title page/preamble)

    if current:
        clauses.append(current)

    return clauses


def sub_split_long_clause(text, max_chars, overlap):
    """Fallback fixed-size splitter, only used when a single clause is too long."""
    if len(text) <= max_chars:
        return [text]
    parts = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        parts.append(text[start:end])
        start = end - overlap if end < len(text) else end
    return parts


def chunk_document(filepath):
    filename = os.path.basename(filepath)
    if filepath.lower().endswith(".pdf"):
        raw_text = read_pdf_file(filepath)
    elif filepath.lower().endswith(".docx"):
        raw_text = read_docx_file(filepath)
    else:
        raw_text = read_text_file(filepath)

    doc_title = filename
    first_line = raw_text.strip().splitlines()[0] if raw_text.strip() else ""
    m = DOC_TITLE_RE.match(first_line.strip())
    if m:
        doc_title = f"TS {m.group(2)} - {m.group(3)}"

    clauses = split_into_clauses(raw_text, filename, doc_title)

    chunks = []
    for clause in clauses:
        body = "\n".join(clause["body_lines"]).strip()
        if not body:
            continue
        sub_parts = sub_split_long_clause(body, MAX_CHUNK_CHARS, CHUNK_OVERLAP_CHARS)
        for i, part in enumerate(sub_parts):
            chunk_id = f"{filename}::{clause['clause_number']}"
            if len(sub_parts) > 1:
                chunk_id += f"::part{i+1}"
            chunks.append({
                "chunk_id": chunk_id,
                "source_doc": doc_title,
                "source_file": filename,
                "clause_number": clause["clause_number"],
                "clause_title": clause["clause_title"],
                "text": f"{clause['clause_number']} {clause['clause_title']}\n{part}",
            })
    return chunks


def run_ingestion():
    os.makedirs(INDEX_DIR, exist_ok=True)
    all_chunks = []
    files = [f for f in os.listdir(DATA_DIR) if f.lower().endswith((".txt", ".pdf", ".docx"))]
    if not files:
        print(f"No .txt, .pdf, or .docx files found in {DATA_DIR}. "
              f"Unzip a downloaded 3GPP spec (e.g. 24501-id0.zip) and drop the .docx file here.")
        return []

    for fname in sorted(files):
        fpath = os.path.join(DATA_DIR, fname)
        chunks = chunk_document(fpath)
        print(f"  {fname}: {len(chunks)} chunks")
        all_chunks.extend(chunks)

    with open(CHUNKS_PATH, "w", encoding="utf-8") as f:
        for c in all_chunks:
            f.write(json.dumps(c) + "\n")

    print(f"\nTotal: {len(all_chunks)} chunks written to {CHUNKS_PATH}")
    return all_chunks


if __name__ == "__main__":
    print(f"Ingesting documents from {DATA_DIR} ...")
    run_ingestion()
    """
Ingestion: turns raw 3GPP documents (PDF, DOCX, or .txt) into structure-aware
chunks.

Note on 3GPP's real distribution format: specs downloaded from
https://www.3gpp.org/ftp/Specs/latest/<Release>/<series>_series are .zip
files (e.g. 24501-id0.zip for TS 24.501). Each zip contains a single .doc or
.docx file — not a PDF. Unzip the archive and drop the .docx into
data/sample_3gpp/; this script reads .docx directly (see read_docx_file).
If you only have the older binary .doc format, convert it to .docx first
(e.g. open and re-save in Word, or use LibreOffice: `soffice --convert-to docx`).

Why structure-aware chunking matters for hallucination reduction:
3GPP specs use rigid hierarchical clause numbering (e.g. "5.5.1.2.3 Initial
registration accepted by the network"). If we split naively by character
count, we routinely cut a clause in half and hand the model incomplete
context, which is a direct cause of hallucinated or fabricated procedure
steps. Instead we detect clause headers with a regex and treat each clause
as an atomic unit, only sub-splitting a clause if it exceeds MAX_CHUNK_CHARS.
Every chunk keeps its clause number + title + source document as metadata,
which lets the generator cite exact clauses instead of vague paraphrases.
"""
import os
import re
import json
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from config import DATA_DIR, CHUNKS_PATH, INDEX_DIR, MAX_CHUNK_CHARS, CHUNK_OVERLAP_CHARS

# Matches lines like "5.5.1.2.3 Initial registration accepted by the network"
CLAUSE_HEADER_RE = re.compile(r"^(\d+(?:\.\d+){0,5})\s+(.+)$")

# Table-of-contents lines look identical to real clause headers (a clause
# number followed by a title) but end in a tab-separated page number, e.g.
# "5.5.1\tRegistration procedure\t50". Real in-body clause headers never end
# this way, so this is a reliable way to filter ToC noise out of the index
# without a separate "skip until body starts" heuristic.
TOC_ENTRY_RE = re.compile(r"\t\d+\s*$")

# Matches "3GPP TS 24.501 - Non-Access-Stratum ..." as the doc title line
DOC_TITLE_RE = re.compile(r"^3GPP\s+(TS|TR)\s+([\d.]+)\s*-\s*(.+)$")


def read_text_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def read_pdf_file(path):
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def read_docx_file(path):
    """3GPP's real distribution format is .docx inside a .zip. python-docx
    reads paragraph text; tables (used for some annexes) are skipped for
    simplicity — extend this if you need annex table content indexed too."""
    import docx
    document = docx.Document(path)
    return "\n".join(p.text for p in document.paragraphs)


def split_into_clauses(raw_text, doc_id, doc_title):
    """Split a document's text into (clause_number, clause_title, body) tuples."""
    lines = raw_text.splitlines()
    clauses = []
    current = None

    for line in lines:
        if TOC_ENTRY_RE.search(line):
            continue  # skip table-of-contents lines entirely — not real clause content
        stripped = line.strip()
        m = CLAUSE_HEADER_RE.match(stripped)
        # Real clause titles start with either a capital letter ("Scope",
        # "References", "General") or a digit — many real 3GPP IE names
        # begin with a number, e.g. "5GS mobile identity", "5G-GUTI",
        # "4G-GUTI". This rejects only the actual false-positive pattern:
        # stray lines from the front-matter "version numbering convention"
        # boilerplate (e.g. "3 or greater indicates TSG approved document
        # under change control."), which start with a bare digit followed
        # by a LOWERCASE word — never a digit followed by an uppercase
        # word like "5GS" or "5G-GUTI". An earlier version of this check
        # used "not isupper()" instead of "islower()", which also (and
        # incorrectly) rejected any title starting with a digit — silently
        # dropping real clauses like "9.11.3.4 5GS mobile identity" from
        # the corpus entirely. That specific miss was traced end-to-end
        # from a hallucinated citation in generation, back through
        # retrieval, to this exact filtering bug.
        if m and m.group(2)[:1].islower():
            m = None
        if m and len(m.group(2)) < 120:  # header lines are short; avoids false positives on body text
            if current:
                clauses.append(current)
            current = {
                "clause_number": m.group(1),
                "clause_title": m.group(2).strip(),
                "body_lines": [],
            }
        else:
            if current is not None:
                current["body_lines"].append(line)
            # lines before the first detected clause header are dropped (title page/preamble)

    if current:
        clauses.append(current)

    return clauses


def sub_split_long_clause(text, max_chars, overlap):
    """Fallback fixed-size splitter, only used when a single clause is too long."""
    if len(text) <= max_chars:
        return [text]
    parts = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        parts.append(text[start:end])
        start = end - overlap if end < len(text) else end
    return parts


def chunk_document(filepath):
    filename = os.path.basename(filepath)
    if filepath.lower().endswith(".pdf"):
        raw_text = read_pdf_file(filepath)
    elif filepath.lower().endswith(".docx"):
        raw_text = read_docx_file(filepath)
    else:
        raw_text = read_text_file(filepath)

    doc_title = filename
    first_line = raw_text.strip().splitlines()[0] if raw_text.strip() else ""
    m = DOC_TITLE_RE.match(first_line.strip())
    if m:
        doc_title = f"TS {m.group(2)} - {m.group(3)}"

    clauses = split_into_clauses(raw_text, filename, doc_title)

    chunks = []
    for clause in clauses:
        body = "\n".join(clause["body_lines"]).strip()
        if not body:
            continue
        sub_parts = sub_split_long_clause(body, MAX_CHUNK_CHARS, CHUNK_OVERLAP_CHARS)
        for i, part in enumerate(sub_parts):
            chunk_id = f"{filename}::{clause['clause_number']}"
            if len(sub_parts) > 1:
                chunk_id += f"::part{i+1}"
            chunks.append({
                "chunk_id": chunk_id,
                "source_doc": doc_title,
                "source_file": filename,
                "clause_number": clause["clause_number"],
                "clause_title": clause["clause_title"],
                "text": f"{clause['clause_number']} {clause['clause_title']}\n{part}",
            })
    return chunks


def run_ingestion():
    os.makedirs(INDEX_DIR, exist_ok=True)
    all_chunks = []
    files = [f for f in os.listdir(DATA_DIR) if f.lower().endswith((".txt", ".pdf", ".docx"))]
    if not files:
        print(f"No .txt, .pdf, or .docx files found in {DATA_DIR}. "
              f"Unzip a downloaded 3GPP spec (e.g. 24501-id0.zip) and drop the .docx file here.")
        return []

    for fname in sorted(files):
        fpath = os.path.join(DATA_DIR, fname)
        chunks = chunk_document(fpath)
        print(f"  {fname}: {len(chunks)} chunks")
        all_chunks.extend(chunks)

    with open(CHUNKS_PATH, "w", encoding="utf-8") as f:
        for c in all_chunks:
            f.write(json.dumps(c) + "\n")

    print(f"\nTotal: {len(all_chunks)} chunks written to {CHUNKS_PATH}")
    return all_chunks


if __name__ == "__main__":
    print(f"Ingesting documents from {DATA_DIR} ...")
    run_ingestion()
